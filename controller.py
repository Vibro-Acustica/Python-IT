from PyQt6.QtCore import pyqtSlot, QObject
from PyQt6.QtWidgets import QFileDialog, QMessageBox, QListWidgetItem
from PyQt6.QtCore import Qt
from DataReader import DWDataReader
from model import Dewesoft, TubeSetupModel, ResultsModel, MeasurementModel, ProcessingModel, DataStore, ReportModel, TestConditionsModel, SamplesModel, DocumentationModel, WarningsModel
import matplotlib.pyplot as plt
from PyQt6.QtGui import QPixmap
from io import BytesIO
from PyQt6.QtCore import Qt, pyqtSignal
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from datetime import date
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from lxml import etree
from docx.oxml import OxmlElement
from datetime import date
import logging
from dotenv import load_dotenv


from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QTabWidget,
                             QLabel, QLineEdit, QPushButton, QGridLayout, QListWidget,
                             QComboBox, QTextBrowser, QMessageBox, QStatusBar, QHBoxLayout, QFrame,
                             QFileDialog, QProgressBar, QCheckBox, QTabWidget, QListWidgetItem, QSizePolicy, QGroupBox, QAbstractItemView, QSpinBox)  # Importação corrigida
from PyQt6.QtGui import QFont, QDoubleValidator, QColor
from PyQt6.QtCore import Qt


class BaseController(QObject):
    def __init__(self):
        super().__init__()
        self.logger = logging.getLogger(self.__class__.__name__)

    def handle_error(self, error, title="Error", show_user=True):
        """Centralized error handling for all controllers."""
        self.logger.error(f"{title}: {str(error)}")
        if show_user:
            QMessageBox.critical(None, title, str(error))

    def validate_input(self, data, rules):
        """Basic input validation."""
        errors = []
        for field, value in data.items():
            if field in rules:
                if 'required' in rules[field] and rules[field]['required'] and not value:
                    errors.append(f"{field} is required")
                if 'type' in rules[field]:
                    try:
                        rules[field]['type'](value)
                    except ValueError:
                        errors.append(f"{field} must be of type {rules[field]['type'].__name__}")
        return errors

class TubeSetupController:
    def __init__(self, model, view):
        self.model = model  # TubeSetupModel, for example
        self.view = view
        # Connect the view to the controller
        self.view.set_controller(self)

    def save_measurements(self):
        try:
            # Save the data to the model
            self.model.set_data(self.view.get_input_data())
            # Show success message
            QMessageBox.information(self.view, "Success", "Tube setup measurements saved successfully!")
        except Exception as e:
            QMessageBox.critical(self.view, "Error", f"Failed to save measurements: {str(e)}")
    
    def reset_fields(self):
        # Reset logic (if necessary)
        self.view.mic_spacing.clear()
        self.view.mic1_sample.clear()
        self.view.mic2_sample.clear()
        self.view.tube_diameter.clear()

class SamplesController:
    def __init__(self, model : SamplesModel, view, signals=None):
        self.model = model  # TubeSetupModel, for example
        self.view = view
        self.signals = signals or SamplesSignalEmitter()
        # Connect the view to the controller
        self.view.set_controller(self)

    def save_samples(self, data):
        try:
            self.model.save_sample(data)
            self.signals.samples_updated.emit()  # Emit signal when sample is saved
            QMessageBox.information(self.view, "Success", "Sample data saved successfully!")
        except Exception as e:
            QMessageBox.critical(self.view, "Error", f"Failed to save sample: {str(e)}")

class ResultsController:
    def __init__(self, model : ResultsModel, view, signals):
        self.model = model
        self.view = view
        self.view.set_controller(self)
        self.selected_metrics = set()
        self.selected_measurement = None

        # Populate the view with existing processed measurements
        self.populate_measurements()

        # Connect metric checkboxes to the toggle function
        self.setup_metric_connections()

        signals.measurement_results_updated.connect(self.populate_measurements)

    def populate_measurements(self):
        """Populates the QListWidget with processed measurements from the model."""
        print("Populated Called")
        self.view.concluded_measurements.clear()
        for measurement in self.model.get_processed_measurements():
            print(measurement)
            item = QListWidgetItem(measurement)
            item.setFlags(item.flags() | Qt.ItemFlag.ItemIsUserCheckable)
            item.setCheckState(Qt.CheckState.Unchecked)
            self.view.concluded_measurements.addItem(item)

    def setup_metric_connections(self):
        """Connects checkboxes to toggle_metric function."""
        self.view.original_signal.stateChanged.connect(lambda state: self.toggle_metric("Original Signal", state))
        self.view.fft_signal_graph.stateChanged.connect(lambda state: self.toggle_metric("Fourier Transform", state))
        self.view.calibration_graph.stateChanged.connect(lambda state: self.toggle_metric("Calibration Function", state))
        self.view.absorption_coef_graph.stateChanged.connect(lambda state: self.toggle_metric("Absorption Coefficient", state))
        self.view.reflection_coef_graph.stateChanged.connect(lambda state: self.toggle_metric("Reflection Coefficient", state))
        self.view.impedance_ratio_graph.stateChanged.connect(lambda state: self.toggle_metric("Impedance Ratio", state))
        self.view.admittance_ratio_graph.stateChanged.connect(lambda state: self.toggle_metric("Admittance Ratio", state))
        self.view.transfer_function_graph.stateChanged.connect(lambda state: self.toggle_metric("Transfer Function", state))
        self.view.impedance_graph.stateChanged.connect(lambda state: self.toggle_metric("Impedance", state))
        self.view.propagation_constant_graph.stateChanged.connect(lambda state: self.toggle_metric("Propagation Constant", state))

    def toggle_metric(self, metric_name, state):
        """Handles the selection/deselection of evaluation metrics and updates the graph."""
        print("metric toggled")
        print(state)
        if state == Qt.CheckState.Checked.value:
            self.selected_metrics.add(metric_name)
            print("checked checkbox")
            print(self.selected_metrics)
        else:
            self.selected_metrics.discard(metric_name)
        
        self.display_graph(metric_name)

    def display_graph(self, name):
        """Gera e exibe um gráfico para as métricas selecionadas."""
        print(f"Display Graph called, for metric {name}")
        if not self.selected_metrics:
            print(f"Selected metrics")
            print(self.selected_metrics)
            self.view.set_graph(QPixmap())  # Limpa o gráfico caso não haja métricas selecionadas
            return
        
        selected_items = self.view.concluded_measurements.selectedItems()
        if len(selected_items) == 0:
            self.selected_measurement = selected_items.text()
        elif len(selected_items) > 1:
            #thow error that says only one measurement can be selected
            raise ValueError("Only one measurement can be selected")
        else:
            self.selected_measurement = selected_items[0].text()

        # Obter dados correspondentes do modelo
        fig = self.model.generate_plot(name,self.selected_measurement)

        # Converter gráfico em QPixmap
        buffer = BytesIO()
        fig.savefig(buffer, format="png")
        buffer.seek(0)
        pixmap = QPixmap()
        pixmap.loadFromData(buffer.getvalue())

        # Atualizar a view com o gráfico
        self.view.set_graph(pixmap)

        # Fechar a figura do Matplotlib para liberar memória
        plt.close(fig)

class MeasurementsSignalEmitter(QObject):
    measurement_results_updated = pyqtSignal()

class SamplesSignalEmitter(QObject):
    samples_updated = pyqtSignal()

class MeasurementController():
    def __init__(self, model : MeasurementModel, view, samples_signals=None):
        self.model = model
        self.view = view
        self.view.set_controller(self)
        load_dotenv()
        dewesoft_setup_path = os.getenv("DEWESOFT_SETUP_PATH")
        dewesoft_files_path = os.getenv("DEWESOFT_READ_FILES_PATH")
        self.dreader = DWDataReader(file_paths=dewesoft_files_path)
        self.dewesoft = Dewesoft()
        self.dewesoft.set_sample_rate(1000) 
        self.dewesoft.set_dimensions(800, 600)
        self.dewesoft.load_setup(dewesoft_setup_path)

        self.signals = MeasurementsSignalEmitter()
        
        # Connect to samples signals if provided
        if samples_signals:
            samples_signals.samples_updated.connect(self.populate_samples)
            
        # Populate the view with initial data
        self.populate_samples()
        self.populate_results()

    def populate_samples(self):
        """Populates the QListWidget with available samples."""
        self.view.amostras_listbox.clear()
        for sample in self.model.get_samples():
            self.view.amostras_listbox.addItem(sample)

    def populate_results(self):
        """Populates the QListWidget with measurement results."""
        self.view.resultados_listbox.clear()
        for result in self.model.get_measurement_results():
            self.view.resultados_listbox.addItem(result)

        self.signals.measurement_results_updated.emit()  # Notify others
        
    def load_both_measurements(self, sample_name, normal_name, switched_name):
        """Load both normal and switched measurements for a sample."""
        # Load normal measurement
        self.dreader.open_data_file(normal_name)
        data_normal = self.dreader.get_measurements_as_dataframe()
        print("Normal measurement read")
        self.model.add_measurement_result(data_normal, normal_name)
        self.dreader.close()

        # Load switched measurement
        self.dreader.open_data_file(switched_name)
        data_switched = self.dreader.get_measurements_as_dataframe()
        print("Switched measurement read")
        self.model.add_measurement_result(data_switched, switched_name)
        self.dreader.close()

        self.populate_results()

    def start_measurement(self, sample_name: str):
        """Handles the measurement process."""
        print("Buttom StartMeasurement Clicked")
        # Ask the user for a file name
        from PyQt6.QtWidgets import QInputDialog
        file_name, ok = QInputDialog.getText(self.view, "Enter File Name It has to be the same as the one in Dewesoft", "Measurement file name:")
        if not ok or not file_name.strip():
            print("Measurement cancelled or no file name provided.")
            return
        file_name = file_name.strip()
        time_measured, ok = QInputDialog.getText(self.view, "Enter the measurement time in seconds", "Measurement time (s):")
        if not ok or not file_name.strip():
            print("Measurement cancelled or time was given")
            return
        time_measured = int(time_measured)

        selected_item = self.view.amostras_listbox.currentItem()
        if selected_item:
            sample_name = selected_item.text()
            full_file_name = f"{file_name}_{sample_name}"
            self.dewesoft.measure(time_measured, full_file_name)
            self.dewesoft.close()

            self.dreader.open_data_file(full_file_name)
            names, num =self.dreader.get_channel_list()
            #print("channel names: ", [name.name.decode('utf-8') for name in names])
            #print("channel number: ",num)
            self.model.save_channel_names([name.name.decode('utf-8') for name in names])
            data = self.dreader.get_measurements_as_dataframe()
            print("measurements read")
            self.model.add_measurement_result(data, file_name)
            self.dreader.close()

            # If you want to add more measurements, repeat with the new file name
            # self.dreader.open_data_file("TestAbsorcao_MicTrocado")
            # data = self.dreader.get_measurements_as_dataframe()
            # print("measurements read")
            # self.model.add_measurement_result(data, "TestAbsorcao_MicTrocado")
            # self.dreader.close()

            self.populate_results()

class TestConditionsController:
    def __init__(self, model, view):
        self.model = model
        self.view = view
        self.view.set_controller(self)

    def save_measurements(self):
        try:
            temp = self.view.temp.text()
            temp_unit = self.view.temp_unit.currentText()
            self.model.set_temperature(temp,temp_unit)
            pressure = self.view.pressure.text()
            pressure_unit = self.view.pressure_unit.currentText()
            self.model.set_pressure(pressure, pressure_unit)
            humidity = self.view.humidity.text()
            self.model.set_humidity(humidity)
            self.model.save_all()
            QMessageBox.information(self.view, "Success", "Test conditions saved successfully!")
        except Exception as e:
            QMessageBox.critical(self.view, "Error", f"Failed to save test conditions: {str(e)}")

class ProcessingController:
    def __init__(self, model, view, measurements_signals):
        self.model = model
        self.view = view
        self.view.set_controller(self)
        
        # Connect signals to slots
        self.connect_signals(measurements_signals)
        
        # Populate the view with available measurements
        self.populate_measurements()
        
        # Initialize options panels
        self.setup_options_panels()
        
    def connect_signals(self, measurement_signals):
        """Connect UI signals to controller methods."""
        # Connect checkbox signals
        self.view.average_checkbox.stateChanged.connect(self.on_operation_changed)
        self.view.combine_checkbox.stateChanged.connect(self.on_operation_changed)
        self.view.extract_checkbox.stateChanged.connect(self.on_operation_changed)
        
        # Connect buttons
        self.view.compute_button.clicked.connect(self.on_compute_clicked)
        self.view.select_all_button.clicked.connect(self.select_all_measurements)
        self.view.deselect_all_button.clicked.connect(self.deselect_all_measurements)
        
        # Connect list widget item changed signal
        self.view.measurements_checklist.itemChanged.connect(self.on_measurement_selection_changed)

        measurement_signals.measurement_results_updated.connect(self.populate_measurements)
    
    def populate_measurements(self):
        """Populate the list with available measurements from the model."""
        self.view.measurements_checklist.clear()
        for measurement in self.model.get_available_measurements():
            item = QListWidgetItem(measurement)
            item.setFlags(item.flags() | Qt.ItemFlag.ItemIsUserCheckable)
            item.setCheckState(Qt.CheckState.Unchecked)
            self.view.measurements_checklist.addItem(item)
    
    def setup_options_panels(self):
        """Create operation-specific option panels."""
        # Clear existing options
        while self.view.options_layout.count():
            item = self.view.options_layout.takeAt(0)
            widget = item.widget()
            if widget:
                widget.deleteLater()
        
        # Create panels for each operation type
        self.average_options = self.create_average_options()
        self.combine_options = self.create_combine_options()
        self.extract_options = self.create_extract_options()
        
        # Hide all panels initially
        self.average_options.setVisible(False)
        self.combine_options.setVisible(False)
        self.extract_options.setVisible(False)
        
        # Add panels to options layout
        self.view.options_layout.addWidget(self.average_options)
        self.view.options_layout.addWidget(self.combine_options)
        self.view.options_layout.addWidget(self.extract_options)
    
    def create_average_options(self):
        """Create options panel for averaging operation."""
        panel = QWidget()
        layout = QVBoxLayout()
        
        layout.addWidget(QLabel("Averaging Options"))
        
        # Add options specific to averaging
        self.average_method = QComboBox()
        self.average_method.addItems(["Arithmetic Mean", "Geometric Mean", "Weighted Average"])
        
        layout.addWidget(QLabel("Method:"))
        layout.addWidget(self.average_method)
        
        panel.setLayout(layout)
        return panel
    
    def create_combine_options(self):
        """Create options panel for combining operation."""
        panel = QWidget()
        layout = QVBoxLayout()
        
        layout.addWidget(QLabel("Combination Options"))
        
        # Add options specific to combining
        self.combine_method = QComboBox()
        self.combine_method.addItems(["Sequential", "Parallel", "Custom"])
        
        layout.addWidget(QLabel("Method:"))
        layout.addWidget(self.combine_method)
        
        panel.setLayout(layout)
        return panel
    
    def create_extract_options(self):
        """Create options panel for third-octave extraction."""
        panel = QWidget()
        layout = QVBoxLayout()
        
        layout.addWidget(QLabel("Third-Octave Options"))
        
        # Add options for third-octave extraction
        self.freq_range_min = QSpinBox()
        self.freq_range_min.setRange(20, 20000)
        self.freq_range_min.setValue(125)
        
        self.freq_range_max = QSpinBox()
        self.freq_range_max.setRange(20, 20000)
        self.freq_range_max.setValue(4000)
        
        freq_layout = QHBoxLayout()
        freq_layout.addWidget(QLabel("Range:"))
        freq_layout.addWidget(self.freq_range_min)
        freq_layout.addWidget(QLabel("Hz to"))
        freq_layout.addWidget(self.freq_range_max)
        freq_layout.addWidget(QLabel("Hz"))
        
        layout.addLayout(freq_layout)
        
        panel.setLayout(layout)
        return panel
    
    def on_operation_changed(self):
        """Handle operation selection changes."""
        # Show/hide option panels based on checkbox selection
        self.view.options_panel.setVisible(True)
        
        self.average_options.setVisible(self.view.average_checkbox.isChecked())
        self.combine_options.setVisible(self.view.combine_checkbox.isChecked())
        self.extract_options.setVisible(self.view.extract_checkbox.isChecked())
        
        # Enable compute button if at least one operation is selected
        has_operations = (self.view.average_checkbox.isChecked() or 
                          self.view.combine_checkbox.isChecked() or 
                          self.view.extract_checkbox.isChecked())
        
        self.view.compute_button.setEnabled(has_operations)
        
        # Hide options panel if no operations selected
        if not has_operations:
            self.view.options_panel.setVisible(False)
    
    def on_compute_clicked(self):
        """Handle compute button click."""
        # Get selected measurements
        selected_measurements = self.get_selected_measurements()
        
        if not selected_measurements:
            QMessageBox.warning(self.view, "No Selection", 
                               "Please select at least one measurement to process.")
            return
        
        # Get selected operations
        operations = []
        if self.view.average_checkbox.isChecked():
            operations.append({
                'type': 'average',
                'method': self.average_method.currentText()
            })
            
        if self.view.combine_checkbox.isChecked():
            operations.append({
                'type': 'combine',
                'method': self.combine_method.currentText()
            })
            
        if self.view.extract_checkbox.isChecked():
            operations.append({
                'type': 'extract',
                'min_freq': self.freq_range_min.value(),
                'max_freq': self.freq_range_max.value()
            })
        
        # Call model method to process the data
        try:
            result = self.model.process_measurements(selected_measurements, operations)
            QMessageBox.information(self.view, "Success", 
                                   f"Successfully processed {len(selected_measurements)} measurements!")
        except Exception as e:
            QMessageBox.critical(self.view, "Processing Error", 
                                f"An error occurred during processing: {str(e)}")
    
    def get_selected_measurements(self):
        """Get list of selected measurements from checklist."""
        selected = []
        for i in range(self.view.measurements_checklist.count()):
            item = self.view.measurements_checklist.item(i)
            if item.checkState() == Qt.CheckState.Checked:
                selected.append(item.text())
        return selected
    
    def select_all_measurements(self):
        """Select all measurements in the list."""
        for i in range(self.view.measurements_checklist.count()):
            item = self.view.measurements_checklist.item(i)
            item.setCheckState(Qt.CheckState.Checked)
    
    def deselect_all_measurements(self):
        """Deselect all measurements in the list."""
        for i in range(self.view.measurements_checklist.count()):
            item = self.view.measurements_checklist.item(i)
            item.setCheckState(Qt.CheckState.Unchecked)
    
    def on_measurement_selection_changed(self, item):
        """Handle changes in measurement selection."""
        # Update UI based on selection if needed
        selected_count = len(self.get_selected_measurements())
        self.view.compute_button.setEnabled(selected_count > 0 and 
                                          (self.view.average_checkbox.isChecked() or 
                                           self.view.combine_checkbox.isChecked() or 
                                           self.view.extract_checkbox.isChecked()))


class ReportGeneratorController(BaseController):
    def __init__(self, model: ReportModel, view):
        super().__init__()
        self.model = model
        self.view = view
        self.view.set_controller(self)
        
        # Load any existing data
        self.load_existing_data()

    def load_existing_data(self):
        """Load existing data into the view."""
        client_info = self.model.get_client_info()
        sample_info = self.model.get_sample_info()
        self.view.set_data(client_info, sample_info)

    def handle_preview(self):
        """Handle preview report generation."""
        try:
            # Update model with current view data
            self._update_model_from_view()
            
            # Validate data
            is_valid, errors = self.model.validate_report_data()
            if not is_valid:
                error_msg = "Please fix the following errors:\n" + "\n".join(errors)
                self.view.show_error(error_msg)
                return
            
            # Generate preview
            self.model.refresh_data()
            self.generate_report("preview_report.docx")
            self.view.update_status("Preview ready: preview_report.docx")
            
            # Open the preview file
            os.startfile("preview_report.docx")
            
        except Exception as e:
            self.handle_error(e, "Failed to generate preview")

    def handle_export(self):
        """Handle final report export."""
        try:
            # Update model with current view data
            self._update_model_from_view()
            
            # Validate data
            is_valid, errors = self.model.validate_report_data()
            if not is_valid:
                error_msg = "Please fix the following errors:\n" + "\n".join(errors)
                self.view.show_error(error_msg)
                return
            
            # Ask user for save location
            file_name, _ = QFileDialog.getSaveFileName(
                self.view,
                "Save Report",
                "",
                "Word Documents (*.docx)"
            )
            
            if file_name:
                if not file_name.endswith('.docx'):
                    file_name += '.docx'
                
                self.model.refresh_data()
                self.generate_report(file_name)
                self.view.update_status(f"Report exported successfully: {file_name}")
                
                # Ask if user wants to open the file
                reply = QMessageBox.question(
                    self.view,
                    'Open Report',
                    'Would you like to open the report?',
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
                )
                
                if reply == QMessageBox.StandardButton.Yes:
                    os.startfile(file_name)
                
        except Exception as e:
            self.handle_error(e, "Failed to export report")

    def _update_model_from_view(self):
        """Update model with current view data."""
        client_info = self.view.get_client_info()
        sample_info = self.view.get_sample_info()
        
        self.model.update_client_info(client_info)
        self.model.update_sample_info(sample_info)

    def generate_report(self, filename="report.docx"):
        """Generate a report with real measurement data."""
        try:
            doc = Document()
            
            # Set up document margins
            self._setup_document_margins(doc)
            
            # Add header with logo and contact info
            self._add_header(doc)
            
            # Add report title and standard reference
            self._add_title_section(doc)
            
            # Add client and test information
            self._add_info_section(doc)
            
            # Add test samples and conditions
            self._add_samples_and_conditions(doc)
            
            # Add measurement details and equipment
            self._add_measurement_details(doc)
            
            # Add results section
            self._add_results_section(doc)
            
            # Save the document
            doc.save(filename)
            
        except Exception as e:
            raise Exception(f"Failed to generate report: {str(e)}")

    def _setup_document_margins(self, doc):
        """Set up document margins."""
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

    def _add_header(self, doc):
        """Add header with logo and contact information."""
        header_table = doc.add_table(rows=3, cols=2)
        header_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        header_table.autofit = False

        # Add logo
        header_cell_1 = header_table.cell(0, 0)
        header_cell_1.merge(header_table.cell(2, 0))
        logo_run = header_cell_1.paragraphs[0].add_run()
        try:
            logo_run.add_picture('logo2.png', width=Inches(2.0))
        except Exception as e:
            self.logger.warning(f"Could not load logo: {str(e)}")

        # Add contact info
        self._add_contact_info(header_table)
        self._remove_table_borders(header_table)

    def _add_contact_info(self, table):
        """Add contact information to header table."""
        # Company info
        info_cell = table.cell(0, 1).paragraphs[0]
        info_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        info_run = info_cell.add_run("VibroAcustica Acoustics Laboratories\nJoinville – SC, 89219-600, Brazil")
        self._format_run(info_run, size=8)

        # Phone
        phone_cell = table.cell(1, 1).paragraphs[0]
        phone_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        phone_run = phone_cell.add_run("+55 (47) 3440-1787")
        self._format_run(phone_run, size=8)

        # Page number
        page_cell = table.cell(2, 1).paragraphs[0]
        page_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        page_run = page_cell.add_run("Page 1 of 9")  # TODO: Implement dynamic page numbering
        self._format_run(page_run, size=8)

    def _add_title_section(self, doc):
        """Add report title and standard reference."""
        # Add spacing
        for _ in range(3):
            doc.add_paragraph()

        # Report title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title.add_run("LABORATORY TEST REPORT")
        self._format_run(title_run, size=12, bold=True, color=RGBColor(220, 0, 0))

        # Standard reference
        standard = doc.add_paragraph()
        standard.alignment = WD_ALIGN_PARAGRAPH.LEFT
        standard_run = standard.add_run(
            "BS EN ISO 10534-2: 2001. DETERMINATION OF SOUND ABSORPTION COEFFICIENT\n"
            "IN IMPEDANCE TUBES. TRANSFER-FUNCTION METHOD."
        )
        self._format_run(standard_run, size=10, bold=True, color=RGBColor(220, 0, 0))

    def _add_info_section(self, doc):
        """Add client and test information section."""
        # Add spacing
        for _ in range(2):
            doc.add_paragraph()

        metadata = self.model.get_report_metadata()
        client_info = metadata['client_info']

        info_table = self._create_info_table(doc, client_info)
        self._remove_table_borders(info_table)

    def _create_info_table(self, doc, client_info):
        """Create and populate the information table with client data."""
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        # Client Information
        self._add_table_row(table, 0, "Client:", client_info.get('name', 'N/A'))
        self._add_table_row(table, 1, "Test Reference:", client_info.get('test_ref', 'N/A'))
        self._add_table_row(table, 2, "Date:", client_info.get('date', str(date.today())))
        self._add_table_row(table, 3, "Report No:", client_info.get('report_no', 'N/A'))
        
        return table

    def _add_table_row(self, table, row, label, value):
        """Add a row to the table with label and value."""
        # Label cell
        label_cell = table.cell(row, 0)
        label_para = label_cell.paragraphs[0]
        label_run = label_para.add_run(label)
        self._format_run(label_run, size=10, bold=True)
        
        # Value cell
        value_cell = table.cell(row, 1)
        value_para = value_cell.paragraphs[0]
        value_run = value_para.add_run(str(value))
        self._format_run(value_run, size=10)

    def _add_samples_and_conditions(self, doc):
        """Add test samples and conditions section."""
        doc.add_page_break()
        
        # Get real data from model
        test_conditions = self.model.report_data['test_conditions']
        sample_info = self.model.report_data['sample_info']
        
        self._add_styled_heading(doc, "1. TEST SAMPLES AND CONDITIONS", 1)
        
        # Samples section
        self._add_styled_heading(doc, "1.1. Description of Test Samples", 2)
        for key, value in sample_info.items():
            self._add_styled_paragraph(doc, f"{key}: {value}")
        
        # Conditions section
        self._add_styled_heading(doc, "1.2. Test Conditions", 2)
        if test_conditions:
            temp = test_conditions.get('temp', {})
            self._add_styled_paragraph(doc, f"Temperature: {temp.get('value', 'N/A')} {temp.get('unit', '°C')}")
            self._add_styled_paragraph(doc, f"Humidity: {test_conditions.get('humi', {}).get('value', 'N/A')} %")
            pressure = test_conditions.get('pressure', {})
            self._add_styled_paragraph(doc, f"Pressure: {pressure.get('value', 'N/A')} {pressure.get('unit', 'hPa')}")

    def _add_measurement_details(self, doc):
        """Add measurement details and equipment section."""
        self._add_styled_heading(doc, "2. MEASUREMENT DETAILS", 1)
        
        # Equipment section
        self._add_styled_heading(doc, "2.1. Equipment", 2)
        equipment_list = self.model.report_data['equipment']
        self._add_equipment_table(doc, equipment_list)
        
        # Procedure section
        self._add_styled_heading(doc, "2.2. Procedure", 2)
        self._add_procedure_content(doc)
        
        # Calculations section
        self._add_styled_heading(doc, "2.3. Calculations", 2)
        self._add_calculations_content(doc)

    def _add_results_section(self, doc):
        """Add results section with real measurement data."""
        self._add_styled_heading(doc, "3. RESULTS", 1)
        
        # Get real absorption data
        absorption_data = self.model.get_absorption_data()
        
        if absorption_data:
            # Create results table
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Add headers
            headers = ["Frequency\nHz", "αₙ (tube Ø100 mm)", "αₙ (tube Ø29 mm)", "αₙ (combined)"]
            for idx, text in enumerate(headers):
                cell = table.cell(0, idx)
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(text)
                self._format_run(run, size=10)
                cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add data rows
            for freq, alpha100, alpha29, combined in absorption_data:
                row_cells = table.add_row().cells
                values = [freq, alpha100, alpha29, combined]
                
                for idx, val in enumerate(values):
                    text = "-" if val == 0.00 else f"{val:.2f}" if isinstance(val, float) else str(val)
                    para = row_cells[idx].paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run(text)
                    self._format_run(run, size=10)
            
            # Add graph if available
            self._add_absorption_graph(doc)
        else:
            self._add_styled_paragraph(doc, "No measurement data available.")

    def _format_run(self, run, size=10, bold=False, color=None):
        """Format a text run with consistent styling."""
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color

    def _remove_table_borders(self, table):
        """Remove borders from a table."""
        for row in table.rows:
            for cell in row.cells:
                cell._element.tcPr.tcBorders = None

    def _add_styled_heading(self, doc, text, level):
        """Add a styled heading to the document."""
        heading = doc.add_paragraph()
        heading_text = heading.add_run(text)
        self._format_run(heading_text, size=12 if level == 1 else 11, bold=True)
        return heading

    def _add_styled_paragraph(self, doc, text):
        """Add a styled paragraph to the document."""
        para = doc.add_paragraph()
        para_text = para.add_run(text)
        self._format_run(para_text, size=10)
        return para

    def _add_equipment_table(self, doc, equipment_list):
        """Add equipment table with real data."""
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header
        for idx, text in enumerate(["Equipment", "Serial/Ref No."]):
            cell = table.cell(0, idx)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(text)
            self._format_run(run, size=10, bold=True, color=RGBColor(255, 255, 255))
            self._set_cell_background(cell, "C00000")
        
        # Data rows
        for equipment in equipment_list:
            row_cells = table.add_row().cells
            for idx, text in enumerate([equipment['name'], equipment['serial']]):
                para = row_cells[idx].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.add_run(text)
                self._format_run(run, size=10)

    def _set_cell_background(self, cell, color):
        """Set background color for a table cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)

    def _add_absorption_graph(self, doc):
        """Add absorption coefficient graph if available."""
        try:
            img_paragraph = doc.add_paragraph()
            img_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            img_run = img_paragraph.add_run()
            img_run.add_picture('ac.png', width=Inches(4.5))
            
            # Add caption
            caption = doc.add_paragraph("Figure 2: Absorption Coefficient vs Frequency")
            caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            caption_run = caption.runs[0]
            self._format_run(caption_run, size=10)
            caption_run.italic = True
        except Exception as e:
            self.logger.warning(f"Could not add absorption graph: {str(e)}")

    def _add_procedure_content(self, doc):
        """Add procedure content to the report."""
        para = doc.add_paragraph()
        para.add_run("Describe the measurement procedure here.").italic = True

    def _add_calculations_content(self, doc):
        """Add calculations content to the report."""
        # Add calculation methods and formulas
        calculations = [
            "The sound absorption coefficient (α) is calculated using the transfer function method according to ISO 10534-2:2001.",
            "The transfer function (H12) between the two microphone positions is measured and used to calculate the reflection coefficient (R).",
            "The absorption coefficient is then calculated as: α = 1 - |R|²",
            "The measurements are performed in two impedance tubes with different diameters (29 mm and 100 mm) to cover the frequency range from 100 Hz to 6.4 kHz."
        ]
        
        for calc in calculations:
            para = doc.add_paragraph()
            para.add_run(calc)
            self._format_run(para.runs[0], size=10)

class DocumentationController(BaseController):
    def __init__(self, model, view):
        super().__init__()
        self.model = model
        self.view = view
        self.view.set_controller(self)
        
        # Load initial documentation
        self.load_documentation()
        
        # Connect signals
        self.view.section_list.currentItemChanged.connect(self.on_section_changed)
        
    def load_documentation(self):
        """Load all documentation sections into the view."""
        try:
            sections = self.model.get_all_sections()
            self.view.clear_sections()
            for section_name, section_data in sections.items():
                self.view.add_section(section_name, section_data['title'])
        except Exception as e:
            self.handle_error(e, "Failed to load documentation")
            
    def on_section_changed(self, current, previous):
        """Handle section selection change."""
        if current:
            try:
                section_name = current.data(Qt.ItemDataRole.UserRole)
                section = self.model.get_documentation_section(section_name)
                if section:
                    self.view.display_section(section['title'], section['content'])
            except Exception as e:
                self.handle_error(e, "Failed to load section")
                
    def update_section(self, section_name, title, content):
        """Update a documentation section."""
        try:
            if self.model.update_section(section_name, title, content):
                self.load_documentation()  # Refresh view
                return True
            else:
                self.handle_error("Section not found", "Update Failed")
                return False
        except Exception as e:
            self.handle_error(e, "Failed to update section")
            return False

class WarningsController(BaseController):
    def __init__(self, model: WarningsModel, view):
        super().__init__()
        self.model = model
        self.view = view
        self.view.set_controller(self)

    def save_warning_settings(self):
        """Save warning settings from view to model."""
        try:
            # Get settings from view
            settings = self.view.get_warning_settings()
            
            # Validate settings
            validation_errors = self._validate_settings(settings)
            if validation_errors:
                self.view.show_error_message("\n".join(validation_errors))
                return
            
            # Save to model
            success = self.model.save_warning_settings(
                snr=float(settings['snr']),
                auto_spec=float(settings['auto_spec']),
                calib=float(settings['calib'])
            )
            
            if success:
                self.view.show_success_message("Warning settings saved successfully!")
            else:
                self.view.show_error_message("Failed to save warning settings")
                
        except Exception as e:
            self.handle_error(e, "Failed to save warning settings")

    def load_warning_settings(self):
        """Load warning settings from model to view."""
        try:
            settings = self.model.get_warning_settings()
            self.view.set_warning_settings(settings)
        except Exception as e:
            self.handle_error(e, "Failed to load warning settings")

    def _validate_settings(self, settings):
        """Validate warning settings."""
        errors = []
        
        # Check for empty values
        for field, value in settings.items():
            if not value.strip():
                errors.append(f"{field.replace('_', ' ').title()} is required")
                continue
            
            # Check numeric values
            try:
                float_val = float(value)
                # Add specific validation rules
                if field == 'snr' and float_val < 0:
                    errors.append("Signal-to-Noise Ratio must be positive")
            except ValueError:
                errors.append(f"{field.replace('_', ' ').title()} must be a valid number")
        
        return errors

class MainAppController:
    def __init__(self, main_app_view):
        self.main_app_view = main_app_view

        self.data_store = DataStore()  # Shared storage
        
        # Create signal emitters
        self.samples_signals = SamplesSignalEmitter()
        
        # Initialize controllers for each tab and connect them to views
        self.tube_setup_controller = self.create_tube_setup_controller()
        self.samples_controller = self.create_samples_controller()
        self.test_conditions_controller = self.create_test_conditions_controller()
        self.measurements_controller = self.create_measurements_controller()
        self.processing_controller = self.create_post_processing_controller(self.measurements_controller.signals)
        self.results_controller = self.create_results_controller(self.measurements_controller.signals)
        self.report_controller = self.create_report_controller()
        self.documentation_controller = self.create_documentation_controller()
        self.warnings_controller = self.create_warnings_controller()

    def create_tube_setup_controller(self):
        tube_setup_model = TubeSetupModel(self.data_store)  # Assuming the TubeSetupModel is defined earlier
        tube_setup_view = self.main_app_view.tube_setup_tab
        controller = TubeSetupController(tube_setup_model, tube_setup_view)
        return controller

    def create_samples_controller(self):
        samples_model = SamplesModel(self.data_store)
        samples_view = self.main_app_view.samples_tab
        controller = SamplesController(samples_model, samples_view, self.samples_signals)
        return controller

    def create_test_conditions_controller(self):
        conditions_model = TestConditionsModel(self.data_store)
        conditions_view = self.main_app_view.test_conditions_tab
        controller = TestConditionsController(conditions_model, conditions_view)
        return controller

    def create_report_controller(self):
        report_model = ReportModel(self.data_store)
        report_view = self.main_app_view.export_tab
        controller = ReportGeneratorController(report_model, report_view)
        return controller

    def create_documentation_controller(self):
        documentation_model = DocumentationModel(self.data_store)
        documentation_view = self.main_app_view.documentation_tab
        controller = DocumentationController(documentation_model, documentation_view)
        return controller

    def create_warnings_controller(self):
        warnings_model = WarningsModel(self.data_store)
        warnings_view = self.main_app_view.warnings_tab
        controller = WarningsController(warnings_model, warnings_view)
        return controller

    def create_measurements_controller(self):
        self.measurement_model = MeasurementModel(self.data_store)
        self.measurement_view = self.main_app_view.measurements_tab.measure_tab
        measurement_controller = MeasurementController(self.measurement_model, self.measurement_view, self.samples_signals)
        return measurement_controller

    def create_results_controller(self, signals):
        self.results_model = ResultsModel(self.data_store)
        self.results_view = self.main_app_view.results_tab
        results_controller = ResultsController(self.results_model, self.results_view, signals)
        return results_controller

    def create_post_processing_controller(self, signals):
        self.post_processing_model = ProcessingModel(self.data_store)
        self.processing_view = self.main_app_view.measurements_tab.processing_tab
        processing_controller = ProcessingController(self.post_processing_model, self.processing_view, signals)
        return processing_controller

    def load_existing_data(self):
        # Load data from models (if necessary)
        pass

